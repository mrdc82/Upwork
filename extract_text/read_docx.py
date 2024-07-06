# import docx NOT python-docx 
from docx import Document 
  
# create an instance of a word document 
doc = Document("example.docx") 

numparagraphs = len(doc.paragraphs)
docxdict = {}

for i in range(numparagraphs):
    x = doc.paragraphs[i].text
    x = x.lower()
    x = x.split(None)
    #print(x)
    for word in x:
        if word not in docxdict:
            docxdict[word] = 0
            docxdict[word] += 1
        else:
            docxdict[word] += 1  

print(docxdict)