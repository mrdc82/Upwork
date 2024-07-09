# import docx NOT python-docx 
from docx import Document
from pypdf import PdfReader
import re
import pandas as pd
from tkinter import Tk
from tkinter.filedialog import askdirectory as ad
import os
import csv

srcdir = ad(title='Select Source Folder') # shows dialog box and return the path
print(srcdir)
#destdir = ad(title='Select Destination Folder') # shows dialog box and return the path
#print(destdir)

try: 
    with open("output.csv", "x") as outputfile:
         outputfile.write("")
except FileExistsError: 
    print(f"The file output.csv exists.")

def docfile():
    docdict = {}

    with open("example.doc","r") as file:
        x = file.read()
        x = x.lower()
        x = re.sub("[()/!.<>:#$,0123456789-]", "", x)
        x = x.split(None)

        for word in x:
            if word not in docdict:
                docdict[word] = 0
                docdict[word] += 1
            else:
                docdict[word] += 1
            with open("output.csv","r") as outputfile:
                if docdict.keys() in outputfile:
                    docdict[word] += docdict.values()

    df = pd.DataFrame(zip(docdict.keys(),docdict.values()), columns=['word','count'])
    df.to_csv("output.csv", index=False)

def docxfile():
# create an instance of a word document 
    doc = Document("example.docx") 

    numparagraphs = len(doc.paragraphs)
    docxdict = {}

    for i in range(numparagraphs):
        x = doc.paragraphs[i].text
        x = x.lower()
        x = re.sub("[()/!.<>:#$,0123456789-]", "", x)
        x = x.split(None)
        #print(x)
        for word in x:
            if word not in docxdict:
                docxdict[word] = 0
                docxdict[word] += 1
            else:
                docxdict[word] += 1
    df = pd.DataFrame(zip(docxdict.keys(),docxdict.values()), columns=['word','count'])
    df.to_csv("output.csv", index=False)

def pdffile():
    pdfile = "example.pdf"
    total = 0
    pdfdict = {}
    # creating a pdf reader object 
    reader = PdfReader(pdfile) 

    # printing number of pages in pdf file 
    numpages = (len(reader.pages)) 
    #print(numpages)

    # creating a page object 
    for i in range(numpages):
        page = reader.pages[i]

    # extracting text from page 
        x = (page.extract_text())
        x = x.lower()
        x = re.sub("[()/!.<>:#$,0123456789-]", "", x)
        paragraph = x.split(None)

        for word in paragraph:
            if word not in pdfdict:
                pdfdict[word] = 0
                pdfdict[word] += 1
            else:
                pdfdict[word] += 1
            with open("output.csv","r") as outputfile:
                if pdfdict.keys() in outputfile:
                    pdfdict[word] += pdfdict.values()

    df = pd.DataFrame(zip(pdfdict.keys(),pdfdict.values()), columns=['word','count'])
    df.to_csv("output.csv", index=False)

def txtfile():
    txtdict = {}
    with open("example.txt","r") as file:
        x = file.read()
        x = x.lower()
        x = re.sub("[()/!.<>:#$,0123456789-]", "", x)
        x = x.split(None)

        for word in x:
            if word not in txtdict:
                txtdict[word] = 0
                txtdict[word] += 1
            else:
                txtdict[word] += 1
    df = pd.DataFrame(zip(txtdict.keys(),txtdict.values()), columns=['word','count'])
    df.to_csv("output.csv", index=False)

if __name__ == '__main__':
    # Iterate over files in directory
    for fname in os.listdir(srcdir):
        if fname.endswith('.doc'):
            docfile()
        elif fname.endswith('.docx'):
            docxfile()
        elif fname.endswith('.pdf'):
            pdffile()
        elif fname.endswith('.txt'):
            txtfile()
        else:
            with open("invalidfiles.txt", "a+") as errorfile:
                errorfile.writelines(fname+"\n")